import streamlit as st
import sqlite3
import bcrypt
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import datetime
import urllib.parse

# ==========================================
# KONFIGURASI HALAMAN & CSS
# ==========================================
st.set_page_config(
    page_title="Monitoring Binpres KONI",
    page_icon="🏆",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS untuk Desain Keren & Profesional
st.markdown("""
    <style>
    .main-header { font-size: 38px; font-weight: 800; color: #1E3A8A; text-align: center; margin-bottom: -10px; }
    .sub-header { font-size: 18px; font-weight: 400; color: #64748B; text-align: center; margin-bottom: 30px; }
    .stButton>button { border-radius: 8px; font-weight: 600; transition: 0.3s; }
    .stButton>button:hover { transform: scale(1.02); }
    .info-box { background-color: #F8FAFC; border-left: 5px solid #3B82F6; padding: 15px; border-radius: 5px; margin-bottom: 20px;}
    .divider { height: 2px; background: linear-gradient(90deg, #1E3A8A 0%, #3B82F6 100%); margin: 20px 0; }
    </style>
""", unsafe_allow_html=True)

# ==========================================
# 1. KONFIGURASI DATABASE & AUTHENTICATION
# ==========================================
DB_NAME = "monitoring.db"

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users (username TEXT PRIMARY KEY, password TEXT, role TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS schedules (id INTEGER PRIMARY KEY AUTOINCREMENT, cabor TEXT, tanggal TEXT, tempat TEXT)''')

    # Buat akun admin default jika belum ada
    c.execute("SELECT * FROM users WHERE username='admin'")
    if not c.fetchone():
        hashed_pw = bcrypt.hashpw('admin123'.encode('utf-8'), bcrypt.gensalt())
        c.execute("INSERT INTO users (username, password, role) VALUES (?, ?, ?)", ('admin', hashed_pw, 'admin'))
    conn.commit()
    conn.close()

def authenticate(username, password):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT password, role FROM users WHERE username=?", (username,))
    result = c.fetchone()
    conn.close()

    if result:
        stored_password, role = result
        if bcrypt.checkpw(password.encode('utf-8'), stored_password):
            return True, role
    return False, None

# Fungsi Format Waktu Bahasa Indonesia
def get_current_time_id():
    now = datetime.datetime.now()
    hari = ["Senin", "Selasa", "Rabu", "Kamis", "Jumat", "Sabtu", "Minggu"]
    bulan = ["Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
    return f"{hari[now.weekday()]}, {now.day} {bulan[now.month - 1]} {now.year} - {now.strftime('%H:%M')} WIB"

# ==========================================
# 2. FUNGSI GENERATE WORD (.docx)
# ==========================================
def generate_word_report(petugas_text, cabor, tanggal, tempat, catatan, fotos):
    doc = Document()
    
    # Header Dokumen
    head = doc.add_heading('LAPORAN MONITORING CABANG OLAHRAGA', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info Jadwal
    doc.add_paragraph(f"Cabang Olahraga\t: {cabor}")
    doc.add_paragraph(f"Tanggal\t\t: {tanggal}")
    doc.add_paragraph(f"Tempat\t\t: {tempat}\n")
    
    # Daftar Petugas
    doc.add_heading('Daftar Petugas:', level=3)
    petugas_list = [p.strip() for p in petugas_text.split('\n') if p.strip()]
    for i, p in enumerate(petugas_list, 1):
        doc.add_paragraph(f"{i}. {p}")

    # Catatan Evaluasi
    doc.add_heading('\nCatatan Evaluasi / Hasil Monitoring:', level=3)
    doc.add_paragraph(catatan)

    # DOKUMENTASI (Pindah ke halaman baru agar rapi di 1 halaman)
    if fotos:
        doc.add_page_break()
        head_doc = doc.add_heading('Lampiran Foto Dokumentasi', level=2)
        head_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Cabor: {cabor} | Tanggal: {tanggal}\n")
        
        # Buat tabel 2 kolom agar muat banyak foto di 1 lembar
        table = doc.add_table(rows=0, cols=2)
        table.autofit = False 
        
        row_cells = None
        for idx, foto in enumerate(fotos):
            if idx % 2 == 0:
                row_cells = table.add_row().cells
            
            cell = row_cells[idx % 2]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            
            try:
                image_stream = io.BytesIO(foto.getvalue())
                run.add_picture(image_stream, width=Inches(2.8)) 
            except Exception as e:
                run.add_text(f"(Gagal memuat gambar: {e})")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 3. HALAMAN & ANTARMUKA PENGGUNA
# ==========================================
init_db()

if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
    st.session_state['username'] = ''
    st.session_state['role'] = ''

# --- HALAMAN LOGIN ---
if not st.session_state['logged_in']:
    st.markdown("<div class='main-header'>🏆 E-MONEV CABOR</div>", unsafe_allow_html=True)
    st.markdown("<div class='sub-header'>Binpres KONI Kabupaten Tangerang</div>", unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1.5, 2, 1.5])
    with col2:
        with st.container(border=True):
            st.markdown("#### 🔐 Silakan Masuk")
            with st.form("login_form"):
                username_input = st.text_input("👤 Username")
                password_input = st.text_input("🔑 Password", type="password")
                submit_btn = st.form_submit_button("Masuk Sistem", use_container_width=True)

                if submit_btn:
                    is_auth, role = authenticate(username_input, password_input)
                    if is_auth:
                        st.session_state['logged_in'] = True
                        st.session_state['username'] = username_input
                        st.session_state['role'] = role
                        st.rerun()
                    else:
                        st.error("🚨 Username atau password salah!")

# --- HALAMAN UTAMA (SETELAH LOGIN) ---
else:
    # Sidebar
    st.sidebar.markdown("### 🏆 PANEL MONEV")
    st.sidebar.caption("Binpres KONI Kab. Tangerang")
    st.sidebar.markdown(f"**🕒 Waktu Sistem:**\n*{get_current_time_id()}*")
    st.sidebar.markdown("---")
    
    st.sidebar.info(f"👤 **Login:** {st.session_state['username'].upper()}\n\n🛡️ **Role:** {st.session_state['role'].upper()}")
    
    # Menu Navigasi Berdasarkan Role
    if st.session_state['role'] == 'admin':
        menu = ["📅 Kelola Jadwal (Admin)", "👥 Kelola User (Admin)", "📝 Coba Isi Laporan"]
        choice = st.sidebar.radio("📌 Navigasi Admin:", menu)
    else:
        choice = "📝 Isi Form Laporan"
        st.sidebar.success("✅ Silakan isi form laporan di panel kanan.")

    st.sidebar.markdown("---")
    if st.sidebar.button("🚪 Keluar (Logout)", use_container_width=True, type="secondary"):
        st.session_state.clear()
        st.rerun()

    # ----------------------------------------------------
    # HALAMAN UNTUK USER: FORM LAPORAN
    # ----------------------------------------------------
    if choice in ["📝 Isi Form Laporan", "📝 Coba Isi Laporan"]:
        st.markdown(f"### 📝 Form Laporan Monitoring")
        st.markdown(f"**Tanggal Hari Ini:** {get_current_time_id()}")
        st.markdown("<div class='divider'></div>", unsafe_allow_html=True)
        
        conn = sqlite3.connect(DB_NAME)
        c = conn.cursor()
        c.execute("SELECT id, cabor, tanggal, tempat FROM schedules")
        schedules_data = c.fetchall()
        conn.close()

        if not schedules_data:
            st.warning("⚠️ Belum ada jadwal monitoring yang tersedia. Harap hubungi Admin.")
        else:
            schedule_options = {f"{s[1]} | {s[2]} | {s[3]}": s for s in schedules_data}
            selected_label = st.selectbox("📌 1. Pilih Jadwal Monitoring yang Tersedia", list(schedule_options.keys()))
            selected_schedule = schedule_options[selected_label]
            
            val_cabor = selected_schedule[1]
            val_tanggal = selected_schedule[2]
            val_tempat = selected_schedule[3]

            with st.container(border=True):
                st.markdown("#### 📋 2. Detail Evaluasi & Dokumentasi")
                
                petugas_text = st.text_area("👤 Daftar Petugas (Tulis 1 nama per baris)", 
                                           placeholder="Contoh:\nBudi Santoso\nAndi Saputra", height=100)
                
                catatan = st.text_area("✍️ Catatan Evaluasi / Hasil Monitoring", height=150)
                
                st.markdown("**📸 Upload Foto Bukti (Bebas 2 s/d 5 Foto)**")
                fotos = st.file_uploader("Otomatis digabung jadi 1 halaman rapi di Word.", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True)

                submit_laporan = st.button("📄 Generate Dokumen Laporan", use_container_width=True, type="primary")

            if submit_laporan:
                if not petugas_text.strip():
                    st.error("⚠️ Harap isi minimal 1 nama petugas!")
                elif not catatan.strip():
                    st.error("⚠️ Catatan evaluasi tidak boleh kosong!")
                elif len(fotos) < 2:
                    st.error("🚨 Minimal unggah 2 foto dokumentasi.")
                elif len(fotos) > 5:
                    st.error("🚨 Maksimal 5 foto dokumentasi agar muat 1 halaman.")
                else:
                    with st.spinner("⏳ Menyusun dokumen laporan..."):
                        word_file = generate_word_report(petugas_text, val_cabor, val_tanggal, val_tempat, catatan, fotos)
                        
                        # Membersihkan string nama file dari karakter ilegal
                        safe_date_name = val_tanggal.replace(" s/d ", "_").replace("-", "").replace("/", "")
                        file_name_doc = f"Monev_{val_cabor.split()[0]}_{safe_date_name}.docx"
                        
                        st.session_state['report_generated'] = True
                        st.session_state['word_file'] = word_file
                        st.session_state['file_name_doc'] = file_name_doc
                        
                        wa_number = "6285691860578"
                        pesan = f"Halo Admin, berikut adalah Laporan Monitoring *{val_cabor}*.\nTanggal: {val_tanggal}\nTempat: {val_tempat}\n\nLaporan telah saya unduh, berikut adalah file lampirannya."
                        wa_link = f"https://wa.me/{wa_number}?text={urllib.parse.quote(pesan)}"
                        st.session_state['wa_link'] = wa_link
            
            if st.session_state.get('report_generated', False):
                st.success("🎉 **Laporan Selesai! Ikuti 2 langkah di bawah ini:**")
                colA, colB = st.columns(2)
                
                with colA:
                    st.download_button(
                        label="📥 1. Unduh Laporan Word (.docx)",
                        data=st.session_state['word_file'],
                        file_name=st.session_state['file_name_doc'],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                with colB:
                    st.link_button("📲 2. Kirim Laporan via WhatsApp Admin", 
                                   st.session_state['wa_link'], 
                                   use_container_width=True)
                    st.caption("*(Pastikan Anda melampirkan/attach file yang baru saja diunduh saat jendela WhatsApp terbuka)*")

    # ----------------------------------------------------
    # MENU ADMIN: KELOLA JADWAL
    # ----------------------------------------------------
    elif choice == "📅 Kelola Jadwal (Admin)":
        st.markdown(f"### 📅 Kelola Jadwal Monitoring")
        st.markdown(f"**Waktu Saat Ini:** {get_current_time_id()}")
        st.markdown("<div class='divider'></div>", unsafe_allow_html=True)
        
        base_cabor = [
            "ANGGAR (IKASI)", "ATLETIK (PASI)", "ANGKAT BESI (PABSI)", "BOLA BASKET (PERBASI)", 
            "BOLA VOLI (PBVSI)", "BULU TANGKIS (PBSI)", "CATUR (PERCASI)", "E. SPORT",
            "KARATE (FORKI)", "PENCAK SILAT (IPSI)", "RENANG (PRSI)", "SEPAK BOLA (Askab-PSSI)", 
            "TAEKWONDO (TI)", "TENIS MEJA (PTMSI)", "TINJU (PERTINA)", "WUSHU (WI)"
        ]
        
        conn = sqlite3.connect(DB_NAME)
        c = conn.cursor()
        c.execute("SELECT DISTINCT cabor FROM schedules")
        existing_cabors = [row[0] for row in c.fetchall()]
        conn.close()
        
        combined_cabor = sorted(list(set(base_cabor + existing_cabors)))
        combined_cabor.append("➕ LAINNYA (Tambah Baru)")
        
        col_form, col_data = st.columns([1, 1.5])
        
        with col_form:
            with st.container(border=True):
                st.subheader("➕ Tambah Jadwal Baru")
                with st.form("form_jadwal"):
                    
                    selected_cabor_option = st.selectbox("Pilih Cabang Olahraga", combined_cabor)
                    
                    if selected_cabor_option == "➕ LAINNYA (Tambah Baru)":
                        custom_cabor = st.text_input("Ketik Nama Cabor Baru", placeholder="Cth: PANAHAN (PERPANI)")
                    else:
                        custom_cabor = "" 
                        
                    # MENGGUNAKAN TUPLE AGAR MENDUKUNG RENTANG TANGGAL
                    new_tanggal = st.date_input(
                        "Tanggal Kegiatan (Bisa pilih satu hari atau rentang hari)", 
                        value=(datetime.date.today(), datetime.date.today())
                    )
                    
                    new_tempat = st.text_input("Tempat / Lokasi")
                    
                    submit_jadwal = st.form_submit_button("Simpan Jadwal", use_container_width=True)
                    
                    if submit_jadwal:
                        final_cabor = custom_cabor.strip().upper() if selected_cabor_option == "➕ LAINNYA (Tambah Baru)" else selected_cabor_option
                        
                        # LOGIKA PEMROSESAN RENTANG TANGGAL
                        if isinstance(new_tanggal, tuple):
                            if len(new_tanggal) == 2:
                                if new_tanggal[0] == new_tanggal[1]:
                                    # Jika hanya klik 1 tanggal
                                    final_tanggal = new_tanggal[0].strftime("%d-%m-%Y")
                                else:
                                    # Jika klik tanggal awal dan akhir
                                    final_tanggal = f"{new_tanggal[0].strftime('%d-%m-%Y')} s/d {new_tanggal[1].strftime('%d-%m-%Y')}"
                            elif len(new_tanggal) == 1:
                                final_tanggal = new_tanggal[0].strftime("%d-%m-%Y")
                            else:
                                final_tanggal = ""
                        else:
                            final_tanggal = new_tanggal.strftime("%d-%m-%Y")
                        
                        if selected_cabor_option == "➕ LAINNYA (Tambah Baru)" and not final_cabor:
                            st.warning("⚠️ Nama Cabang Olahraga baru tidak boleh kosong!")
                        elif not new_tempat:
                            st.warning("⚠️ Tempat/Lokasi tidak boleh kosong!")
                        elif not final_tanggal:
                            st.warning("⚠️ Tanggal kegiatan tidak boleh kosong! (Pilih dua kali untuk rentang)")
                        else:
                            conn = sqlite3.connect(DB_NAME)
                            c = conn.cursor()
                            c.execute("INSERT INTO schedules (cabor, tanggal, tempat) VALUES (?, ?, ?)", 
                                      (final_cabor, final_tanggal, new_tempat))
                            conn.commit()
                            conn.close()
                            st.success(f"✅ Jadwal {final_cabor} ditambahkan!")
                            st.rerun()
                        
        with col_data:
            with st.container(border=True):
                st.subheader("📋 Daftar Jadwal Aktif")
                conn = sqlite3.connect(DB_NAME)
                jadwal_df = pd.read_sql_query("SELECT id as ID, cabor as Cabor, tanggal as Tanggal, tempat as Tempat FROM schedules", conn)
                conn.close()
                
                if jadwal_df.empty:
                    st.info("Belum ada jadwal yang dibuat.")
                else:
                    st.dataframe(jadwal_df, use_container_width=True, hide_index=True)
                    with st.expander("🗑️ Hapus Jadwal"):
                        del_id = st.selectbox("Pilih ID Jadwal yang akan dihapus", jadwal_df['ID'].tolist())
                        if st.button("Hapus Jadwal", type="primary"):
                            conn = sqlite3.connect(DB_NAME)
                            c = conn.cursor()
                            c.execute("DELETE FROM schedules WHERE id=?", (del_id,))
                            conn.commit()
                            conn.close()
                            st.success("Jadwal berhasil dihapus!")
                            st.rerun()

    # ----------------------------------------------------
    # MENU ADMIN: KELOLA USER
    # ----------------------------------------------------
    elif choice == "👥 Kelola User (Admin)":
        st.markdown(f"### 👥 Manajemen Pengguna")
        st.markdown("<div class='divider'></div>", unsafe_allow_html=True)
        
        tab1, tab2 = st.tabs(["📋 Daftar Pengguna", "➕ Tambah Pengguna Baru"])
        with tab1:
            conn = sqlite3.connect(DB_NAME)
            users_df = pd.read_sql_query("SELECT username as Username, role as 'Hak Akses' FROM users", conn)
            conn.close()
            st.dataframe(users_df, use_container_width=True, hide_index=True)

            with st.expander("🗑️ Hapus Pengguna", expanded=False):
                del_user = st.selectbox("Pilih pengguna yang akan dihapus", users_df['Username'].tolist())
                if st.button("Hapus Akun", type="primary"):
                    if del_user == 'admin':
                        st.error("⚠️ Tidak bisa menghapus akun admin utama!")
                    else:
                        conn = sqlite3.connect(DB_NAME)
                        c = conn.cursor()
                        c.execute("DELETE FROM users WHERE username=?", (del_user,))
                        conn.commit()
                        conn.close()
                        st.success(f"✅ User **{del_user}** berhasil dihapus!")
                        st.rerun()

        with tab2:
            with st.container(border=True):
                with st.form("add_user_form"):
                    new_username = st.text_input("👤 Username Baru")
                    new_password = st.text_input("🔑 Password Baru", type="password")
                    new_role = st.selectbox("🛡️ Hak Akses", ["user", "admin"])
                    submit_new_user = st.form_submit_button("💾 Simpan Pengguna", use_container_width=True)

                    if submit_new_user:
                        if new_username and new_password:
                            try:
                                hashed_pw = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt())
                                conn = sqlite3.connect(DB_NAME)
                                c = conn.cursor()
                                c.execute("INSERT INTO users (username, password, role) VALUES (?, ?, ?)", 
                                          (new_username, hashed_pw, new_role))
                                conn.commit()
                                conn.close()
                                st.success(f"✅ Pengguna baru **{new_username}** berhasil ditambahkan!")
                                st.rerun()
                            except sqlite3.IntegrityError:
                                st.error("⚠️ Username sudah terdaftar!")
                        else:
                            st.warning("⚠️ Username dan Password tidak boleh kosong!")
